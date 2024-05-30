from dash import Dash, dcc, html, Input, Output, State
import dash_bootstrap_components as dbc
import pandas as pd
from sqlalchemy import create_engine
import re
import io
import xlsxwriter
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# إعداد الاتصال بقاعدة البيانات باستخدام SQLAlchemy
connection_string = "mssql+pyodbc://PDASYNC:PDASYNC@52.174.138.11/SBDBLIVE2?driver=ODBC+Driver+18+for+SQL+Server&TrustServerCertificate=yes"
engine = create_engine(connection_string)

# جلب البيانات من قاعدة البيانات
query_warehouse = """
SELECT [BUID],
      [WarehouseName],
      [ItemID],
      [StockByCarton],
      [TotalSalesByCarton],
      [DaysOfStock],
      [BrandNameE],
      [MasterBrandName],
      [ItemNameE],
      [BUIDName]
FROM [SBDBLIVE2].[dbo].[StockTurnoverViewbyBUID];
"""

query_sales = """
SELECT [WAREHOUSEID],
      [WarehouseName],
      [ITEMID],
      [StockByCarton],
      [AvgSalesByCarton],
      [DaysOfStock],
      [BrandNameE],
      [MasterBrandName],
      [ItemNameE],
      [BUIDName]
FROM [SBDBLIVE2].[dbo].[StockTurnoverView];
"""

df_warehouse = pd.read_sql(query_warehouse, engine)
df_sales = pd.read_sql(query_sales, engine)

# قائمة بأسماء المندوبين والفروع التي نريد إزالتها
salesman_names = ['هيثم فاروق فولي', 'اسم مندوب آخر', 'اسم مندوب آخر 2']
branch_names = df_warehouse['BUIDName'].unique().tolist()

def remove_names(item_name, names_list):
    for name in names_list:
        # استخدام التعبير المنتظم لإزالة الأسماء بشكل دقيق
        item_name = re.sub(r'\b' + re.escape(name) + r'\b', '', item_name).strip()
    return item_name

# معالجة عمود ItemNameE لإزالة أسماء المندوبين والفروع
df_warehouse['ItemNameE'] = df_warehouse['ItemNameE'].apply(lambda x: remove_names(x, salesman_names + branch_names))
df_sales['ItemNameE'] = df_sales['ItemNameE'].apply(lambda x: remove_names(x, salesman_names + branch_names))

# إعداد واجهة المستخدم باستخدام Dash
app = Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])

def generate_progress_bar(row):
    days_of_stock = row['DaysOfStock']
    if days_of_stock <= 2:
        color = 'danger'
    elif days_of_stock <= 5:
        color = 'warning'
    else:
        color = 'success'
    return html.Div([
        html.H6(f"{row['ItemNameE']} in {row['WarehouseName']}"),
        dbc.Progress(value=days_of_stock, max=10, color=color, className="mb-3", label=f"{days_of_stock} days")
    ])

app.layout = dbc.Container([
    dbc.Row([
        dbc.Col(html.H1("Warehouse and Sales Stock Dashboard"), className="mb-2")
    ]),
    dbc.Row([
        dbc.Col(html.H6("Analyze and manage warehouse and sales stocks effectively."), className="mb-4")
    ]),
    dbc.Row([
        dbc.Col(html.H5("Stock Status in Sites"), className="mb-2 mt-4")
    ]),
    dbc.Row([
        dbc.Col(dcc.Dropdown(
            id='buid-dropdown',
            options=[{'label': name, 'value': name} for name in df_warehouse['BUIDName'].unique()],
            placeholder="Select a BUID Name",
        ), width=12)
    ]),
    dbc.Row(id='warehouse-progress-bars-site', className="mt-4"),
    dbc.Row([
        dbc.Col(html.H5("Stock Status by Salesman"), className="mb-2 mt-4")
    ]),
    dbc.Row([
        dbc.Col(dcc.Dropdown(
            id='warehouseid-dropdown',
            options=[],  # سيتم تحديث الخيارات بناءً على BUID المختار
            placeholder="Select a Warehouse Name",
        ), width=12)
    ]),
    dbc.Row(id='warehouse-progress-bars-salesman', className="mt-4"),
    dbc.Row([
        dbc.Col(html.H5("Download Data"), className="mb-2 mt-4")
    ]),
    dbc.Row([
        dbc.Col(dbc.Button("Download as Excel", id="btn-download-excel", color="primary"), width=4),
        dbc.Col(dbc.Button("Download as Word", id="btn-download-word", color="primary"), width=4),
        dbc.Col(dbc.Button("Download as PDF", id="btn-download-pdf", color="primary"), width=4)
    ]),
    dcc.Download(id="download-dataframe-xlsx"),
    dcc.Download(id="download-dataframe-docx"),
    dcc.Download(id="download-dataframe-pdf")
], fluid=True)

@app.callback(
    Output('warehouse-progress-bars-site', 'children'),
    Input('buid-dropdown', 'value')
)
def update_warehouse_progress_bars_site(selected_buid_name):
    if selected_buid_name is None:
        return []
    filtered_df = df_warehouse[df_warehouse['BUIDName'] == selected_buid_name]
    return [generate_progress_bar(row) for _, row in filtered_df.iterrows()]

@app.callback(
    Output('warehouseid-dropdown', 'options'),
    Input('buid-dropdown', 'value')
)
def update_salesman_dropdown(selected_buid_name):
    if selected_buid_name is None:
        return []
    filtered_sales = df_sales[df_sales['BUIDName'] == selected_buid_name]
    return [{'label': name, 'value': name} for name in filtered_sales['WarehouseName'].unique()]

@app.callback(
    Output('warehouse-progress-bars-salesman', 'children'),
    Input('warehouseid-dropdown', 'value'),
    State('buid-dropdown', 'value')
)
def update_warehouse_progress_bars_salesman(selected_warehouse_name, selected_buid_name):
    if selected_warehouse_name is None or selected_buid_name is None:
        return []
    filtered_df = df_sales[(df_sales['BUIDName'] == selected_buid_name) & (df_sales['WarehouseName'] == selected_warehouse_name)]
    return [generate_progress_bar(row) for _, row in filtered_df.iterrows()]

@app.callback(
    Output("download-dataframe-xlsx", "data"),
    Input("btn-download-excel", "n_clicks"),
    State('buid-dropdown', 'value'),
    prevent_initial_call=True,
)
def download_excel(n_clicks, selected_buid_name):
    if selected_buid_name is None:
        return
    filtered_df = df_warehouse[df_warehouse['BUIDName'] == selected_buid_name]
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    filtered_df.to_excel(writer, sheet_name='Sheet1', index=False)
    writer.close()
    output.seek(0)
    return dcc.send_bytes(output.getvalue(), "data.xlsx")

@app.callback(
    Output("download-dataframe-docx", "data"),
    Input("btn-download-word", "n_clicks"),
    State('buid-dropdown', 'value'),
    prevent_initial_call=True,
)
def download_word(n_clicks, selected_buid_name):
    if selected_buid_name is None:
        return
    filtered_df = df_warehouse[df_warehouse['BUIDName'] == selected_buid_name]
    document = Document()
    document.add_heading('Filtered Data', 0)
    for index, row in filtered_df.iterrows():
        document.add_paragraph(f"{row['ItemNameE']} in {row['WarehouseName']} with {row['DaysOfStock']} days of stock")
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return dcc.send_bytes(output.getvalue(), "data.docx")

@app.callback(
    Output("download-dataframe-pdf", "data"),
    Input("btn-download-pdf", "n_clicks"),
    State('buid-dropdown', 'value'),
    prevent_initial_call=True,
)
def download_pdf(n_clicks, selected_buid_name):
    if selected_buid_name is None:
        return
    filtered_df = df_warehouse[df_warehouse['BUIDName'] == selected_buid_name]
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    y = height - 40
    for index, row in filtered_df.iterrows():
        text = f"{row['ItemNameE']} in {row['WarehouseName']} with {row['DaysOfStock']} days of stock"
        c.drawString(30, y, text)
        y -= 20
        if y < 40:
            c.showPage()
            y = height - 40
    c.save()
    output.seek(0)
    return dcc.send_bytes(output.getvalue(), "data.pdf")

if __name__ == '__main__':
    app.run_server(debug=True, use_reloader=False)
